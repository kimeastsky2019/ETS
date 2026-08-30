"""
app.py — FastAPI 메인 애플리케이션
[MIGRATION] xAI/Grok → Exo 1.0 + Qdrant + CrewAI

변경 사항:
- xai_sdk, xai_helpers 제거
- mgmt_client (Grok Collections) → Qdrant QdrantClient
- chat_client (xAI) → openai.AsyncOpenAI (Exo 호환)
- /analyze → Exo(Qwen 72B) 직접 호출
- /agent 엔드포인트 신규 추가 (CrewAI)
- Collection.xai_id 필드 = Qdrant 컬렉션 이름으로 재사용
- Document.xai_doc_id 필드 = 인제스트 작업 UUID로 재사용
"""
import asyncio
import secrets
import uuid
import time
import os
import json
import re
import io
from typing import Optional
from contextlib import asynccontextmanager
from datetime import timedelta

import httpx
from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Form, status, BackgroundTasks
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from sqlmodel import select
from sqlalchemy import delete, func
from sqlmodel.ext.asyncio.session import AsyncSession
from openai import AsyncOpenAI
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, OptimizersConfigDiff

from config import (
    EXO_ENABLED, EXO_BASE_URL, EXO_API_KEY, LLM_MODEL,
    INTERNAL_API_TOKEN,
    QDRANT_HOST, QDRANT_PORT, QDRANT_API_KEY,
    EMBED_DIM, CACHE_TTL_SEC, CACHE_MAXSIZE,
    COST_PER_1M_INPUT, COST_PER_1M_OUTPUT,
    SYSTEM_GUARDRAIL,
)
from llm_router import router as llm_router
from cache import cache_get, cache_set
from rag import run_rag
from embeddings import embedding_info
from ingest import (
    ingest_bytes, ensure_collection, collection_name_sanitize,
    delete_document_chunks,
)
from database import init_db, get_session
from models import Collection, Document, User, UsageEvent
import contract
import storage
from auth_utils import (
    verify_password, get_password_hash,
    create_access_token, ACCESS_TOKEN_EXPIRE_MINUTES,
)

# ──────────────────────────────────────────────
# Qdrant 동기 클라이언트 (컬렉션 관리용)
# ──────────────────────────────────────────────
def get_qdrant() -> QdrantClient:
    return QdrantClient(
        host=QDRANT_HOST,
        port=QDRANT_PORT,
        api_key=QDRANT_API_KEY,
    )

# ──────────────────────────────────────────────
# Exo LLM 클라이언트 (분석 엔드포인트용)
# ──────────────────────────────────────────────
def get_exo_client() -> AsyncOpenAI:
    return AsyncOpenAI(base_url=EXO_BASE_URL, api_key=EXO_API_KEY)


# ──────────────────────────────────────────────
# App lifecycle
# ──────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()

    # 기본 유저 생성 (최초 실행 시)
    # 공개 배포에서는 ADMIN_EMAIL/ADMIN_PASSWORD 로 반드시 덮어쓰세요.
    admin_email    = os.getenv("ADMIN_EMAIL", "info@gngmeta.com")
    admin_password = os.getenv("ADMIN_PASSWORD", "admin1234")
    async for session in get_session():
        statement = select(User).where(User.email == admin_email)
        results = await session.exec(statement)
        user = results.first()
        if not user:
            default_user = User(
                email=admin_email,
                hashed_password=get_password_hash(admin_password),
                full_name=os.getenv("ADMIN_NAME", "GnG Admin"),
                acl="RESTRICTED",   # 어드민은 전 등급 열람
            )
            session.add(default_user)
            await session.commit()
        break

    yield


app = FastAPI(title="RAG API — Exo + Qdrant + CrewAI", lifespan=lifespan)

# CORS
allowed_origins = os.getenv("CORS_ORIGINS", "*").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"] if "*" in allowed_origins else allowed_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

# --- 지식 데이터베이스 구축 (업종별 온톨로지 · 규제 준수) ------------------
# 의존성(pdfplumber/openpyxl)이 없으면 본체는 계속 동작한다. 지식DB 화면만 죽는다.
try:
    from kb.router import router as kb_router
    app.include_router(kb_router)
except Exception as _e:
    print(f"Warning: knowledge-base router not loaded: {_e}")


oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")


# ──────────────────────────────────────────────
# Auth helpers
# ──────────────────────────────────────────────
async def get_current_user(
    token: str = Depends(oauth2_scheme),
    session: AsyncSession = Depends(get_session),
):
    from jose import JWTError, jwt
    from auth_utils import SECRET_KEY, ALGORITHM

    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception

    statement = select(User).where(User.email == email)
    result = await session.exec(statement)
    user = result.first()
    if user is None:
        raise credentials_exception
    return user


# ──────────────────────────────────────────────
# Pydantic Models
# ──────────────────────────────────────────────
class Filters(BaseModel):
    category: str | None = None
    tags: list[str] | None = None
    version: str | None = None
    date_from: str | None = None
    date_to: str | None = None
    # 지식 데이터베이스 축. sector 는 업종(닫힌 집합), channel 은 글/표/그림.
    # 표만 검색하면 수치 질의의 정확도가 크게 오른다.
    sector: str | None = None
    channel: str | None = None


class ChatRequest(BaseModel):
    query: str = Field(..., min_length=1)
    # 0 이면 전체 컬렉션을 한 번에 검색한다. 보고서를 사업장별로 나눈 뒤
    # "화학 분야 보고서" 같은 가로지르는 질문을 할 수 없던 것을 푸는 값이다.
    collection_id: int = Field(..., description="검색할 컬렉션 ID (0=전체)")
    filters: Filters | None = None


class ChatResponse(BaseModel):
    request_id: str
    answer: str
    citations: list[dict] = []
    cached: bool
    latency_ms: int
    retrieval: str = ""      # hybrid_rrf | dense — 어떤 검색 경로가 쓰였는지


class AgentRequest(BaseModel):
    query: str = Field(..., min_length=1)
    collection_id: int = Field(..., description="검색할 컬렉션 ID")
    mode: str = Field(
        default="rag_crew",
        description="에이전트 모드: single | rag_crew | research_crew",
    )


class AgentResponse(BaseModel):
    request_id: str
    answer: str
    mode: str
    latency_ms: int


class CollectionCreate(BaseModel):
    name: str
    description: str | None = None
    category: str | None = None
    tags: str | None = None


class CollectionUpdate(BaseModel):
    name: str | None = None
    description: str | None = None
    category: str | None = None
    tags: str | None = None


class CollectionRead(BaseModel):
    id: int
    name: str
    xai_id: str          # Qdrant 컬렉션 이름 저장
    description: str | None = None
    category: str | None = None
    tags: str | None = None
    created_at: str
    documents_count: int | None = None
    processing_count: int | None = None
    failed_count: int | None = None
    status: str | None = None


class DocumentRead(BaseModel):
    id: int
    name: str
    xai_doc_id: str      # 인제스트 작업 UUID
    status: str
    created_at: str
    # ── 데이터 계약 ──
    stable_id: Optional[str] = None
    version: int = 1
    sha256: Optional[str] = None
    acl: str = "INTERNAL"
    owner: Optional[str] = None
    doc_status: str = "draft"
    mime: Optional[str] = None
    size_bytes: Optional[int] = None
    chunk_count: int = 0
    has_original: bool = False


class DocumentDetail(DocumentRead):
    provenance: dict = {}
    collection_id: Optional[int] = None
    citation: str = ""


def _doc_read(d: Document) -> DocumentRead:
    return DocumentRead(
        id=d.id,
        name=d.name,
        xai_doc_id=d.xai_doc_id,
        status=d.status,
        created_at=d.created_at.isoformat(),
        stable_id=d.stable_id,
        version=d.version or 1,
        sha256=d.sha256,
        acl=contract.normalize_acl(d.acl),
        owner=d.owner,
        doc_status=contract.normalize_doc_status(d.doc_status),
        mime=d.mime,
        size_bytes=d.size_bytes,
        chunk_count=d.chunk_count or 0,
        has_original=storage.exists(d.file_path),
    )


class Token(BaseModel):
    access_token: str
    token_type: str


class UserCreate(BaseModel):
    email: str
    password: str
    full_name: Optional[str] = None


class UserRead(BaseModel):
    id: int
    email: str
    full_name: Optional[str] = None
    acl: str = "PUBLIC"          # 이 사용자가 열람 가능한 최대 등급


# ──────────────────────────────────────────────
# Auth Endpoints
# ──────────────────────────────────────────────
@app.post("/token", response_model=Token)
async def login_for_access_token(
    form_data: OAuth2PasswordRequestForm = Depends(),
    session: AsyncSession = Depends(get_session),
):
    statement = select(User).where(User.email == form_data.username)
    result = await session.exec(statement)
    user = result.first()

    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )

    access_token = create_access_token(
        data={"sub": user.email},
        expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES),
    )
    return {"access_token": access_token, "token_type": "bearer"}


@app.post("/register", response_model=UserRead)
async def register(
    user_in: UserCreate,
    session: AsyncSession = Depends(get_session),
):
    statement = select(User).where(User.email == user_in.email)
    result = await session.exec(statement)
    if result.first():
        raise HTTPException(status_code=400, detail="Email already registered")

    user = User(
        email=user_in.email,
        hashed_password=get_password_hash(user_in.password),
        full_name=user_in.full_name,
    )
    session.add(user)
    await session.commit()
    await session.refresh(user)
    return UserRead(id=user.id, email=user.email, full_name=user.full_name)


@app.get("/users/me", response_model=UserRead)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return UserRead(
        id=current_user.id,
        email=current_user.email,
        full_name=current_user.full_name,
        acl=contract.normalize_acl(current_user.acl),
    )


# ──────────────────────────────────────────────
# Health — Circuit Breaker 상태 포함
# ──────────────────────────────────────────────
@app.get("/health")
async def health():
    qdrant_ok = False
    exo_ok = False
    try:
        qdrant = get_qdrant()
        qdrant.get_collections()
        qdrant_ok = True
    except Exception:
        pass
    # 로컬 추론 엔드포인트 점검.
    # 예전에는 EXO_BASE_URL.rstrip('/v1') + '/health' 를 찔렀는데 두 군데가 틀렸습니다:
    #   1) str.rstrip 은 접미사가 아니라 문자 집합 {'/','v','1'} 을 깎습니다.
    #      "http://host:1/v1" → "http://host:" 가 되어 포트가 비고 80(nginx)에 갔습니다.
    #   2) 이 자리에는 Exo 대신 Ollama 가 오는 경우가 많은데 Ollama 에는 /health 가
    #      없어 404 → 항상 down 으로 보고됐습니다(정상 동작 중인데도).
    # 그래서 OpenAI 호환 규격의 /v1/models 로 점검합니다 — Exo·Ollama 둘 다 구현합니다.
    if not EXO_ENABLED:
        exo_ok = None   # 비활성 — "down" 이 아니라 "disabled" 로 보고
    else:
        try:
            base = EXO_BASE_URL.rstrip("/")
            if not base.endswith("/v1"):
                base = f"{base}/v1"
            async with httpx.AsyncClient(timeout=5) as c:
                r = await c.get(f"{base}/models")
                exo_ok = r.status_code == 200
        except Exception:
            exo_ok = False

    # Circuit Breaker 현황
    cb_status = llm_router.status()

    return {
        "ok":     True,
        "qdrant": "up" if qdrant_ok else "down",
        "exo":    "disabled" if exo_ok is None else ("up" if exo_ok else "down"),
        "llm":    cb_status,   # active_provider, providers 별 state
        "embedding": embedding_info(),
    }


# ──────────────────────────────────────────────
# Circuit Breaker 수동 초기화 (관리자용)
# ──────────────────────────────────────────────
class ResetRequest(BaseModel):
    provider: Optional[str] = None  # None → 전체 초기화


@app.post("/admin/circuit-reset")
async def circuit_reset(
    body: ResetRequest,
    current_user: User = Depends(get_current_user),
):
    """Circuit Breaker 수동 초기화. provider 생략 시 전체 초기화."""
    llm_router.reset(body.provider)
    return {
        "reset": body.provider or "all",
        "status": llm_router.status(),
    }


# ──────────────────────────────────────────────
# Stats
# ──────────────────────────────────────────────
@app.get("/stats")
async def get_stats(
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    total_collections = (await session.exec(select(func.count(Collection.id)))).one()
    total_documents   = (await session.exec(select(func.count(Document.id)))).one()
    total_users       = (await session.exec(select(func.count(User.id)))).one()
    total_queries     = (await session.exec(select(func.count(UsageEvent.id)))).one()
    avg_latency       = (await session.exec(select(func.avg(UsageEvent.latency_ms)))).one() or 0
    total_cost        = (await session.exec(select(func.sum(UsageEvent.cost_usd)))).one() or 0

    return {
        "collections":  total_collections,
        "documents":    total_documents,
        "users":        total_users,
        "queries":      total_queries,
        "avg_latency_ms": int(avg_latency),
        "cost_usd":     float(total_cost),
        "stack":        f"Exo/{LLM_MODEL} + Qdrant + CrewAI",
    }


# ──────────────────────────────────────────────
# 내부 API — 서비스 간 호출 전용 (사용자 JWT 가 아니라 공유 토큰으로 인증)
#
# LLM Wiki(work.ets0404.com)가 "RAG 에 이미 있는 문서" 를 골라 위키로 넘길 때 쓴다.
# 사람이 파일을 다시 업로드하지 않아도 되게 하려는 것이고, 같은 호스트의 루프백
# 호출이라 토큰이 브라우저로 나가지 않는다.
# ──────────────────────────────────────────────
from fastapi import Header
from fastapi.responses import FileResponse


def _require_internal(x_internal_token: str | None = Header(default=None)):
    if not INTERNAL_API_TOKEN:
        raise HTTPException(503, "내부 API 가 설정되지 않았습니다 (INTERNAL_API_TOKEN).")
    # 길이가 달라도 상수시간 비교를 쓴다 — 토큰 추측에 응답시간 단서를 주지 않는다.
    if not x_internal_token or not secrets.compare_digest(x_internal_token, INTERNAL_API_TOKEN):
        raise HTTPException(401, "내부 토큰이 유효하지 않습니다.")
    return True


@app.get("/internal/documents")
async def internal_documents(
    _: bool = Depends(_require_internal),
    session: AsyncSession = Depends(get_session),
):
    """위키로 넘길 수 있는 문서 = 원본 파일이 보관된 것만."""
    rows = (await session.exec(
        select(Document, Collection.name)
        .join(Collection, Document.collection_id == Collection.id, isouter=True)
        .where(Document.file_path.is_not(None))
        .order_by(Document.id.desc())
    )).all()
    out = []
    for d, cname in rows:
        # DB 에는 있는데 파일이 사라진 경우를 목록에서 걸러 낸다 — 고를 수 있는데
        # 누르면 404 나는 항목이 화면에 남지 않게.
        if not d.file_path or not os.path.exists(d.file_path):
            continue
        out.append({
            "id":              d.id,
            "name":            d.name,
            "collection_name": cname,
            "stable_id":       d.stable_id,
            "sha256":          d.sha256,
            "size_bytes":      d.size_bytes,
            "chunk_count":     d.chunk_count,
            "acl":             d.acl,
            "created_at":      d.created_at,
        })
    return {"documents": out, "count": len(out)}


@app.get("/internal/documents/{document_id}/file")
async def internal_document_file(
    document_id: int,
    _: bool = Depends(_require_internal),
    session: AsyncSession = Depends(get_session),
):
    """보관된 원본 파일을 그대로 내려준다."""
    d = await session.get(Document, document_id)
    if not d:
        raise HTTPException(404, "문서를 찾을 수 없습니다.")
    if not d.file_path or not os.path.exists(d.file_path):
        raise HTTPException(404, "원본 파일이 보관되어 있지 않습니다.")
    return FileResponse(
        d.file_path,
        media_type=d.mime or "application/pdf",
        filename=d.name,
    )


# ──────────────────────────────────────────────
# 통계 드릴다운 — 대시보드 카드의 숫자를 눌렀을 때 근거를 보여준다.
# /stats 는 개수만 주므로 "6이 어느 문서인지" 를 확인할 방법이 없었다.
# ──────────────────────────────────────────────
@app.get("/stats/documents")
async def stats_documents(
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """'처리된 문서' 카드의 내역 — 전 컬렉션의 문서 목록."""
    rows = (await session.exec(
        select(Document, Collection.name)
        .join(Collection, Document.collection_id == Collection.id, isouter=True)
        .order_by(Document.id.desc())
    )).all()
    return [
        {
            "id":              d.id,
            "name":            d.name,
            "collection_id":   d.collection_id,
            "collection_name": cname,
            "status":          d.status,
            "chunk_count":     d.chunk_count,
            "size_bytes":      d.size_bytes,
            "acl":             d.acl,
            "doc_status":      d.doc_status,
            "owner":           d.owner,
            "created_at":      d.created_at,
        }
        for d, cname in rows
    ]


@app.get("/stats/queries")
async def stats_queries(
    limit: int = 100,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """'검색 쿼리' · '응답 시간' 카드의 내역 — 최근 호출 기록."""
    limit = max(1, min(limit, 500))
    rows = (await session.exec(
        select(UsageEvent, Collection.name)
        .join(Collection, UsageEvent.collection_id == Collection.id, isouter=True)
        .order_by(UsageEvent.id.desc())
        .limit(limit)
    )).all()
    return [
        {
            "id":              e.id,
            "endpoint":        e.endpoint,
            "model":           e.model,
            "collection_id":   e.collection_id,
            "collection_name": cname,
            "latency_ms":      e.latency_ms,
            "total_tokens":    e.total_tokens,
            "cached":          e.cached,
            "created_at":      e.created_at,
        }
        for e, cname in rows
    ]


@app.get("/stats/collections")
async def stats_collections(
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """'활성 컬렉션' 카드의 내역 — 컬렉션별 문서/청크 수."""
    cols = (await session.exec(select(Collection).order_by(Collection.id))).all()
    out = []
    for c in cols:
        docs = (await session.exec(
            select(Document).where(Document.collection_id == c.id)
        )).all()
        out.append({
            "id":          c.id,
            "name":        c.name,
            "qdrant_name": c.xai_id,
            "description": c.description,
            "documents":   len(docs),
            "chunks":      sum(d.chunk_count or 0 for d in docs),
            "processed":   sum(1 for d in docs if d.status == "processed"),
            "created_at":  c.created_at,
        })
    return out


# ──────────────────────────────────────────────
# Collection Endpoints
# ──────────────────────────────────────────────
@app.get("/collections", response_model=list[CollectionRead])
async def list_collections(
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    result = await session.exec(select(Collection))
    collections = result.all()
    output: list[CollectionRead] = []
    for c in collections:
        total_docs = (await session.exec(
            select(func.count(Document.id)).where(Document.collection_id == c.id)
        )).one()
        processing = (await session.exec(
            select(func.count(Document.id)).where(
                (Document.collection_id == c.id) & (Document.status == "processing")
            )
        )).one()
        failed = (await session.exec(
            select(func.count(Document.id)).where(
                (Document.collection_id == c.id) & (Document.status == "failed")
            )
        )).one()
        output.append(CollectionRead(
            id=c.id,
            name=c.name,
            xai_id=c.xai_id,
            description=c.description,
            category=c.category,
            tags=c.tags,
            created_at=c.created_at.isoformat(),
            documents_count=total_docs,
            processing_count=processing,
            failed_count=failed,
            status="processing" if processing > 0 else "active",
        ))
    return output


@app.post("/collections", response_model=CollectionRead)
async def create_collection(
    collection: CollectionCreate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    # Qdrant 컬렉션 이름 생성 (정규화)
    qdrant_name = collection_name_sanitize(collection.name)

    # 중복 이름 방지
    existing = (await session.exec(
        select(Collection).where(Collection.xai_id == qdrant_name)
    )).first()
    if existing:
        qdrant_name = f"{qdrant_name}_{uuid.uuid4().hex[:6]}"

    # Qdrant에 컬렉션 생성
    try:
        qdrant = get_qdrant()
        ensure_collection(qdrant, qdrant_name)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Qdrant Error: {str(e)}")

    # DB에 저장
    db_collection = Collection(
        name=collection.name,
        xai_id=qdrant_name,           # ← Qdrant 컬렉션 이름 저장
        description=collection.description,
        category=collection.category,
        tags=collection.tags,
    )
    session.add(db_collection)
    await session.commit()
    await session.refresh(db_collection)

    return CollectionRead(
        id=db_collection.id,
        name=db_collection.name,
        xai_id=db_collection.xai_id,
        description=db_collection.description,
        category=db_collection.category,
        tags=db_collection.tags,
        created_at=db_collection.created_at.isoformat(),
    )


@app.put("/collections/{collection_id}", response_model=CollectionRead)
async def update_collection(
    collection_id: int,
    body: CollectionUpdate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")

    if body.name is not None:
        collection.name = body.name.strip()
    if body.description is not None:
        collection.description = body.description.strip() or None
    if body.category is not None:
        collection.category = body.category.strip() or None
    if body.tags is not None:
        collection.tags = body.tags.strip() or None

    session.add(collection)
    await session.commit()
    await session.refresh(collection)

    total_docs = (await session.exec(
        select(func.count(Document.id)).where(Document.collection_id == collection.id)
    )).one()

    return CollectionRead(
        id=collection.id,
        name=collection.name,
        xai_id=collection.xai_id,
        description=collection.description,
        category=collection.category,
        tags=collection.tags,
        created_at=collection.created_at.isoformat(),
        documents_count=total_docs,
    )


@app.delete("/collections/{collection_id}")
async def delete_collection(
    collection_id: int,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")

    # Qdrant 컬렉션 삭제 시도
    try:
        qdrant = get_qdrant()
        qdrant.delete_collection(collection.xai_id)
    except Exception as e:
        print(f"Warning: Qdrant 컬렉션 삭제 실패: {e}")

    # DB 삭제
    await session.exec(delete(Document).where(Document.collection_id == collection_id))
    await session.exec(delete(Collection).where(Collection.id == collection_id))
    await session.commit()

    return {"status": "deleted", "id": collection_id}


# ──────────────────────────────────────────────
# Document Endpoints
# ──────────────────────────────────────────────
@app.get("/collections/{collection_id}", response_model=list[DocumentRead])
async def get_collection_documents(
    collection_id: int,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")

    statement = select(Document).where(Document.collection_id == collection_id)
    results = await session.exec(statement)
    documents = results.all()

    # 권한 밖 문서는 목록에서도 뺀다 — 존재 자체가 정보다.
    return [
        _doc_read(d) for d in documents
        if contract.can_read(current_user.acl, d.acl)
    ]


# ──────────────────────────────────────────────
# 문서 레지스트리 (Data Contract API)
#
# 다른 서비스(LLM Wiki 등)는 파일을 복사해 가는 대신 stable_id 로 참조한다.
# 경로가 /registry 인 이유: /docs 는 FastAPI 가 Swagger UI 로 이미 쓰고 있다.
# 캐시가 필요하면 sha256 이 같을 때만 재사용하면 된다.
# ──────────────────────────────────────────────
async def _docs_for_stable_id(session: AsyncSession, stable_id: str) -> list[Document]:
    rows = await session.exec(
        select(Document).where(Document.stable_id == stable_id)
        .order_by(Document.version.desc())
    )
    return list(rows.all())


def _pick_version(docs: list[Document], version: Optional[int]) -> Optional[Document]:
    if not docs:
        return None
    if version is not None:
        return next((d for d in docs if (d.version or 1) == version), None)
    # 버전 미지정이면 reviewed 최신본을 우선한다. 없으면 그냥 최신본.
    reviewed = [d for d in docs if contract.normalize_doc_status(d.doc_status) == "reviewed"]
    return (reviewed or docs)[0]


@app.get("/registry", response_model=list[DocumentRead])
async def list_docs(
    status: Optional[str] = None,          # draft/reviewed/deprecated
    collection_id: Optional[int] = None,
    latest_only: bool = True,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """소비 서비스용 문서 목록. 요청자 등급으로 걸러진다."""
    stmt = select(Document).where(Document.stable_id.is_not(None))
    if collection_id is not None:
        stmt = stmt.where(Document.collection_id == collection_id)
    rows = await session.exec(stmt.order_by(Document.stable_id, Document.version.desc()))
    docs = [d for d in rows.all() if contract.can_read(current_user.acl, d.acl)]

    if status:
        want = contract.normalize_doc_status(status)
        docs = [d for d in docs if contract.normalize_doc_status(d.doc_status) == want]

    if latest_only:
        seen: set[str] = set()
        latest = []
        for d in docs:                      # 이미 version 내림차순
            if d.stable_id in seen:
                continue
            seen.add(d.stable_id)
            latest.append(d)
        docs = latest

    return [_doc_read(d) for d in docs]


@app.get("/registry/{stable_id}", response_model=DocumentDetail)
async def get_doc(
    stable_id: str,
    version: Optional[int] = None,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """
    문서 메타데이터. version 을 주면 그 시점을 고정해서 가져온다
    (`/registry/{stable_id}?version=3`).
    """
    docs = await _docs_for_stable_id(session, stable_id)
    doc = _pick_version(docs, version)
    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다")
    if not contract.can_read(current_user.acl, doc.acl):
        # 권한이 없으면 존재 여부도 알리지 않는다.
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다")

    base = _doc_read(doc)
    return DocumentDetail(
        **base.model_dump(),
        provenance=contract.parse_provenance(doc.provenance),
        collection_id=doc.collection_id,
        citation=contract.format_citation(doc.stable_id, doc.version or 1, doc.sha256),
    )


@app.get("/registry/{stable_id}/versions", response_model=list[DocumentRead])
async def list_doc_versions(
    stable_id: str,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    docs = await _docs_for_stable_id(session, stable_id)
    docs = [d for d in docs if contract.can_read(current_user.acl, d.acl)]
    if not docs:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다")
    return [_doc_read(d) for d in docs]


@app.get("/registry/{stable_id}/file")
async def download_doc(
    stable_id: str,
    version: Optional[int] = None,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """보관된 원본 파일 다운로드. ACL 검사 후 해시를 헤더로 함께 준다."""
    from fastapi.responses import Response

    docs = await _docs_for_stable_id(session, stable_id)
    doc = _pick_version(docs, version)
    if not doc or not contract.can_read(current_user.acl, doc.acl):
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다")

    content = storage.read_original(doc.file_path)
    if content is None:
        raise HTTPException(
            status_code=410,
            detail="원본이 보관되어 있지 않습니다 (계약 도입 이전 업로드).",
        )

    actual = contract.sha256_bytes(content)
    if doc.sha256 and actual != doc.sha256:
        # 대장과 파일이 어긋났다. 조용히 내려주면 변조를 못 잡는다.
        raise HTTPException(status_code=409, detail="무결성 오류: 원본 해시가 대장과 다릅니다")

    from urllib.parse import quote
    return Response(
        content=content,
        media_type=doc.mime or "application/octet-stream",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(doc.name)}",
            "X-Doc-Stable-Id": doc.stable_id or "",
            "X-Doc-Version": str(doc.version or 1),
            "X-Doc-Sha256": actual,
            "X-Doc-Acl": contract.normalize_acl(doc.acl),
        },
    )


class DocStatusUpdate(BaseModel):
    doc_status: Optional[str] = None      # draft/reviewed/deprecated
    acl: Optional[str] = None
    owner: Optional[str] = None


@app.patch("/registry/{stable_id}", response_model=DocumentRead)
async def update_doc_governance(
    stable_id: str,
    body: DocStatusUpdate,
    version: Optional[int] = None,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """
    거버넌스 필드(리뷰 상태·ACL·담당)만 수정한다.
    내용을 바꾸려면 업로드로 새 버전을 만들어야 한다 — 그게 버전의 의미다.
    """
    docs = await _docs_for_stable_id(session, stable_id)
    doc = _pick_version(docs, version)
    if not doc or not contract.can_read(current_user.acl, doc.acl):
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다")

    if body.doc_status is not None:
        doc.doc_status = contract.normalize_doc_status(body.doc_status)
    if body.acl is not None:
        new_acl = contract.normalize_acl(body.acl)
        if not contract.can_read(current_user.acl, new_acl):
            raise HTTPException(status_code=403, detail="본인 등급보다 높게 올릴 수 없습니다")
        doc.acl = new_acl
    if body.owner is not None:
        doc.owner = body.owner

    session.add(doc)
    await session.commit()
    await session.refresh(doc)

    # 청크 payload 의 acl/doc_status 도 맞춰 둔다 (검색 필터가 이 값을 본다).
    try:
        collection = await session.get(Collection, doc.collection_id)
        if collection:
            from qdrant_client.models import Filter, FieldCondition, MatchValue
            get_qdrant().set_payload(
                collection_name=collection.xai_id,
                payload={
                    "acl": contract.normalize_acl(doc.acl),
                    "doc_status": contract.normalize_doc_status(doc.doc_status),
                    "owner": doc.owner or "",
                },
                points=Filter(must=[
                    FieldCondition(key="stable_id", match=MatchValue(value=stable_id)),
                    FieldCondition(key="version", match=MatchValue(value=doc.version or 1)),
                ]),
            )
    except Exception as e:
        print(f"Warning: 청크 payload 동기화 실패: {e}")

    return _doc_read(doc)


@app.delete("/documents/{document_id}")
async def delete_document(
    document_id: int,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    doc = await session.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not contract.can_read(current_user.acl, doc.acl):
        raise HTTPException(status_code=404, detail="Document not found")

    # Qdrant에서 해당 청크 삭제
    try:
        collection = await session.get(Collection, doc.collection_id)
        if collection:
            qdrant = get_qdrant()
            from qdrant_client.models import Filter, FieldCondition, MatchValue
            if doc.stable_id:
                # 계약 문서는 (stable_id, version) 으로 정확히 지운다.
                # 파일명은 버전마다 같을 수 있어 source 로 지우면 다른 버전까지 날아간다.
                delete_document_chunks(
                    qdrant, collection.xai_id, doc.stable_id, doc.version or 1
                )
            else:
                qdrant.delete(
                    collection_name=collection.xai_id,
                    points_selector=Filter(
                        must=[FieldCondition(key="source", match=MatchValue(value=doc.name))]
                    ),
                )
    except Exception as e:
        print(f"Warning: Qdrant 청크 삭제 실패: {e}")

    # 보관된 원본도 함께 지운다 (대장과 파일이 어긋나지 않도록).
    if doc.stable_id and doc.file_path:
        try:
            storage.delete_version(doc.stable_id, doc.version or 1)
        except Exception as e:
            print(f"Warning: 원본 삭제 실패: {e}")

    await session.delete(doc)
    await session.commit()
    return {"status": "deleted", "id": document_id}


MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".doc"}

# provenance 에 남길 파서 이름. 파서를 교체하면 이 값으로 재인제스트 대상을 찾는다.
_PARSER_BY_EXT = {
    ".pdf":  "PyPDF2",
    ".docx": "python-docx",
    ".doc":  "python-docx",
    ".txt":  "plain",
    ".md":   "plain",
}


def _parser_for(ext: str) -> str:
    return _PARSER_BY_EXT.get(ext.lower(), "plain")


def _extract_text_for_analyze(content: bytes, filename: str) -> str:
    """분석용 텍스트 추출 (간단 버전)"""
    ext = os.path.splitext(filename.lower())[1]
    if ext in (".txt", ".md"):
        return content.decode("utf-8", errors="replace")
    if ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            return content.decode("utf-8", errors="replace")
    if ext == ".pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return f"[PDF: {filename}]"
    return content.decode("utf-8", errors="replace")


async def _next_stable_id(session: AsyncSession, slug_source: str) -> str:
    """
    같은 (slug, 연도) 안에서 다음 일련번호를 붙여 stable_id 를 만든다.
    시각을 넣지 않는 이유는 contract.make_stable_id 주석 참고.
    """
    from datetime import datetime, timezone

    year = datetime.now(timezone.utc).year
    prefix = f"doc:ets:{contract.slugify(slug_source)}-{year}-"
    rows = await session.exec(
        select(Document.stable_id).where(Document.stable_id.like(f"{prefix}%"))
    )
    seq = 0
    for sid in rows.all():
        try:
            seq = max(seq, int((sid or "").rsplit("-", 1)[1]))
        except (IndexError, ValueError):
            continue
    return f"{prefix}{seq + 1:03d}"


@app.post("/collections/{collection_id}/upload")
async def upload_document(
    collection_id: int,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
    version: Optional[str] = Form(None),
    date: Optional[str] = Form(None),
    relatedDocs: Optional[str] = Form(None),
    relationship_note: Optional[str] = Form(None),
    policy_note: Optional[str] = Form(None),
    # ── 데이터 계약 입력 ──
    stable_id: Optional[str] = Form(None),   # 주면 그 문서의 새 버전으로 등록
    acl: Optional[str] = Form(None),
    owner: Optional[str] = Form(None),
    doc_status: Optional[str] = Form(None),
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    # 컬렉션 확인
    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")

    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    file_ext = os.path.splitext(file.filename.lower())[1]
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"지원하지 않는 파일 형식입니다. 지원: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="파일 크기 초과 (최대 100MB)")
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="빈 파일입니다")

    # 메타데이터 빌드
    extra_meta: dict = {}
    if version:
        extra_meta["version"] = version
    if date:
        extra_meta["date"] = date
    if relatedDocs:
        extra_meta["related_docs"] = relatedDocs
    if relationship_note:
        extra_meta["relationship_note"] = relationship_note
    if policy_note:
        extra_meta["policy_note"] = policy_note

    tag_list = [t.strip() for t in tags.split(",")] if tags else None

    # ── 데이터 계약: stable_id 발급 또는 새 버전 ─────────────
    doc_sha256 = contract.sha256_bytes(content)

    if stable_id:
        if not contract.is_valid_stable_id(stable_id):
            raise HTTPException(
                status_code=400,
                detail="stable_id 형식이 올바르지 않습니다 (예: doc:ets:audit-2026-031)",
            )
        prev_rows = await session.exec(
            select(Document).where(Document.stable_id == stable_id)
            .order_by(Document.version.desc())
        )
        prev = prev_rows.first()
        if prev is None:
            raise HTTPException(status_code=404, detail=f"알 수 없는 stable_id: {stable_id}")
        if not contract.can_read(current_user.acl, prev.acl):
            raise HTTPException(status_code=403, detail="이 문서를 수정할 권한이 없습니다")
        if prev.sha256 == doc_sha256:
            # 내용이 같으면 버전을 올리지 않는다. 버전은 "내용이 달라졌다"는 신호여야 한다.
            return {
                "status": "unchanged",
                "document_id": prev.id,
                "stable_id": stable_id,
                "version": prev.version,
                "message": "동일한 내용입니다 — 새 버전을 만들지 않았습니다.",
            }
        doc_version = (prev.version or 1) + 1
        doc_acl = contract.normalize_acl(acl or prev.acl)
        doc_owner = owner or prev.owner
    else:
        stable_id = await _next_stable_id(session, os.path.splitext(file.filename)[0])
        doc_version = 1
        doc_acl = contract.normalize_acl(acl)
        doc_owner = owner or current_user.email

    # 업로더가 자기 등급보다 높은 문서를 만들면 스스로 못 읽는다 — 미리 막는다.
    if not contract.can_read(current_user.acl, doc_acl):
        raise HTTPException(
            status_code=403,
            detail=f"본인 열람 등급({contract.normalize_acl(current_user.acl)})보다 "
                   f"높은 등급({doc_acl})으로 등록할 수 없습니다",
        )

    # ── 원본 보관 (계약의 전제) ──────────────────────────
    try:
        storage.ensure_root()
        saved_path, doc_sha256 = storage.save_original(
            stable_id, doc_version, file.filename, content
        )
    except OSError as e:
        raise HTTPException(status_code=500, detail=f"원본 저장 실패: {e}")

    prov = contract.build_provenance(
        original_name=file.filename,
        parser=_parser_for(file_ext),
        source_path=f"upload:{current_user.email}",
        extra={"collection": collection.name, "uploaded_by": current_user.email},
    )

    job_id = str(uuid.uuid4())
    doc = Document(
        name=file.filename,
        xai_doc_id=job_id,
        collection_id=collection.id,
        status="processing",
        stable_id=stable_id,
        version=doc_version,
        sha256=doc_sha256,
        acl=doc_acl,
        owner=doc_owner,
        doc_status=contract.normalize_doc_status(doc_status),
        provenance=prov,
        file_path=saved_path,
        mime=file.content_type or "application/octet-stream",
        size_bytes=len(content),
    )
    session.add(doc)
    await session.commit()
    await session.refresh(doc)

    # 백그라운드에서 Qdrant 인제스트
    qdrant_name = collection.xai_id
    doc_id = doc.id
    file_content = content
    sid, ver, dacl, downer = stable_id, doc_version, doc_acl, doc_owner
    dstatus = doc.doc_status

    async def do_ingest():
        chunks = 0
        try:
            qdrant = get_qdrant()
            ensure_collection(qdrant, qdrant_name)
            # 같은 버전을 재인제스트하는 경우 이전 청크를 먼저 지운다.
            if ver > 1:
                try:
                    delete_document_chunks(qdrant, qdrant_name, sid, ver)
                except Exception:
                    pass
            chunks = ingest_bytes(
                client=qdrant,
                collection_name=qdrant_name,
                content=file_content,
                filename=file.filename,
                category=category,
                tags=tag_list,
                extra_metadata=extra_meta,
                stable_id=sid,
                version=ver,
                sha256=doc_sha256,
                acl=dacl,
                owner=downer,
                doc_status=dstatus,
            )
            if chunks > 0:
                new_status = "processed"
            elif dacl in contract.ACL_NO_INDEX:
                # RESTRICTED 는 의도적으로 색인하지 않는다 — 실패가 아니다.
                new_status = "stored_not_indexed"
            else:
                new_status = "failed"
        except Exception as e:
            print(f"인제스트 실패 ({file.filename}): {e}")
            new_status = "failed"

        async for s in get_session():
            d = await s.get(Document, doc_id)
            if d:
                d.status = new_status
                d.chunk_count = chunks
                s.add(d)
                await s.commit()
            break

    background_tasks.add_task(do_ingest)

    return {
        "status": "processing",
        "document_id": doc.id,
        "job_id": job_id,
        "stable_id": stable_id,
        "version": doc_version,
        "sha256": doc_sha256,
        "acl": doc_acl,
        "citation": contract.format_citation(stable_id, doc_version, doc_sha256),
        "message": "백그라운드에서 인제스트 중입니다. /collections/{id} 에서 상태를 확인하세요.",
    }


# ──────────────────────────────────────────────
# AI Analyze (온톨로지 메타데이터 추천)
# ──────────────────────────────────────────────
ANALYZE_SYSTEM_PROMPT = """당신은 문서 온톨로지 구축을 돕는 전문 AI 어시스턴트입니다.
반드시 아래 JSON 형식으로만 응답하세요:
{
  "category": "문서 카테고리 (정책/재무/기술/법률/인사/연구 등)",
  "tags": ["태그1", "태그2", "태그3"],
  "summary": "문서 내용 2-3줄 요약",
  "consulting": "이 문서를 온톨로지에 통합할 때 고려할 점과 추천사항"
}"""


@app.post("/analyze")
async def analyze_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """문서를 AI(Qwen 2.5 72B)로 분석하여 온톨로지 메타데이터를 추천합니다."""
    content = await file.read()
    filename = file.filename or "unknown"
    text = _extract_text_for_analyze(content, filename)
    if len(text) > 8000:
        text = text[:8000] + "\n...(이하 생략)"

    try:
        client = get_exo_client()
        completion = await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": ANALYZE_SYSTEM_PROMPT},
                {"role": "user",   "content": f"파일명: {filename}\n\n내용:\n{text}"},
            ],
            temperature=0.3,
            max_tokens=512,
        )
        answer = completion.choices[0].message.content.strip()

        # JSON 파싱 (```json ... ``` 블록 처리)
        cleaned = answer
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```[a-z]*\n?", "", cleaned)
            cleaned = re.sub(r"\n?```$", "", cleaned)
        result = json.loads(cleaned)
        return {
            "category":   result.get("category", ""),
            "tags":       result.get("tags", []),
            "summary":    result.get("summary", ""),
            "consulting": result.get("consulting", ""),
        }
    except json.JSONDecodeError:
        return {"category": "", "tags": [], "summary": "", "consulting": answer}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 분석 실패: {str(e)}")


COLLECTION_ANALYZE_PROMPT = """당신은 문서 컬렉션 온톨로지 구축을 돕는 전문 AI 어시스턴트입니다.
반드시 아래 JSON 형식으로만 응답하세요:
{
  "description": "컬렉션 설명 (2-3줄)",
  "category": "컬렉션 카테고리",
  "tags": "태그1, 태그2, 태그3",
  "consulting": "이 컬렉션을 온톨로지에 통합할 때 고려할 점"
}"""


@app.post("/collections/{collection_id}/analyze")
async def analyze_collection(
    collection_id: int,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """컬렉션의 문서 목록을 AI로 분석하여 메타데이터를 추천합니다."""
    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")

    statement = select(Document).where(Document.collection_id == collection_id)
    results = await session.exec(statement)
    documents = results.all()

    doc_list = "\n".join(f"- {d.name} (상태: {d.status})" for d in documents) if documents else "(문서 없음)"
    user_msg = (
        f"컬렉션 이름: {collection.name}\n"
        f"현재 설명: {collection.description or '(없음)'}\n"
        f"현재 카테고리: {collection.category or '(없음)'}\n"
        f"현재 태그: {collection.tags or '(없음)'}\n"
        f"문서 수: {len(documents)}\n\n"
        f"포함된 문서 목록:\n{doc_list}"
    )

    try:
        client = get_exo_client()
        completion = await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": COLLECTION_ANALYZE_PROMPT},
                {"role": "user",   "content": user_msg},
            ],
            temperature=0.3,
            max_tokens=512,
        )
        answer = completion.choices[0].message.content.strip()
        cleaned = answer
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```[a-z]*\n?", "", cleaned)
            cleaned = re.sub(r"\n?```$", "", cleaned)
        result = json.loads(cleaned)
        return {
            "description": result.get("description", ""),
            "category":    result.get("category", ""),
            "tags":        result.get("tags", ""),
            "consulting":  result.get("consulting", ""),
        }
    except json.JSONDecodeError:
        return {"description": "", "category": "", "tags": "", "consulting": answer}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 분석 실패: {str(e)}")


# ──────────────────────────────────────────────
# Chat Endpoint (RAG)
# ──────────────────────────────────────────────
@app.post("/chat", response_model=ChatResponse)
async def chat(
    req: ChatRequest,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    request_id = str(uuid.uuid4())
    t0 = time.time()

    # collection_id=0 → 전체 검색. 컬렉션을 하나 고르게 하면 여러 사업장을
    # 가로지르는 질문("화학 분야 보고서")을 아예 할 수 없다.
    search_all = req.collection_id == 0
    if search_all:
        rows = (await session.exec(
            select(Collection.id, Collection.xai_id, Collection.name)
        )).all()
        counts = {
            cid: n for cid, n in (await session.exec(
                select(Document.collection_id, func.count(Document.id))
                .where(Document.status == "processed")
                .group_by(Document.collection_id)
            )).all()
        }
        usable = [(cid, xid, name) for cid, xid, name in rows if counts.get(cid)]
        if not usable:
            return ChatResponse(
                request_id=request_id,
                answer="인덱싱된 문서가 없습니다. 먼저 문서를 업로드하고 인제스트를 기다려 주세요.",
                citations=[], cached=False,
                latency_ms=int((time.time() - t0) * 1000),
            )
        qdrant_names = [xid for _, xid, _ in usable]
        qdrant_collection = "|".join(sorted(qdrant_names))   # 캐시 키 용도
        db_collection = None
        doc_count = sum(counts.get(cid, 0) for cid, _, _ in usable)
    else:
        db_collection = await session.get(Collection, req.collection_id)
        if not db_collection:
            raise HTTPException(status_code=404, detail="지정한 컬렉션을 찾을 수 없습니다.")
        qdrant_collection = db_collection.xai_id
        qdrant_names = None
        doc_count = (await session.exec(
            select(func.count(Document.id)).where(
                (Document.collection_id == db_collection.id) & (Document.status == "processed")
            )
        )).one()
    if doc_count == 0:
        return ChatResponse(
            request_id=request_id,
            answer="인덱싱된 문서가 없습니다. 먼저 문서를 업로드하고 인제스트를 기다려 주세요.",
            citations=[],
            cached=False,
            latency_ms=int((time.time() - t0) * 1000),
        )

    filters_dict = req.filters.model_dump(exclude_none=True) if req.filters else None

    # 캐시 확인 — 열람 등급이 키에 포함되어야 등급 간 답변이 섞이지 않는다.
    viewer_acl = contract.normalize_acl(current_user.acl)
    cached = cache_get(qdrant_collection, LLM_MODEL, req.query, filters_dict, viewer_acl)
    if cached:
        return ChatResponse(
            request_id=request_id,
            answer=cached["answer"],
            citations=cached.get("citations", []),
            cached=True,
            latency_ms=int((time.time() - t0) * 1000),
            retrieval=cached.get("retrieval", ""),
        )

    # RAG 실행
    # 모든 LLM 프로바이더가 죽었거나(회로 개방) API 키가 비어 있으면 라우터가
    # RuntimeError 를 던집니다. 그대로 두면 프론트에 맨 500 만 떨어져서
    # "왜 안 되는지" 를 알 수 없으므로 원인을 실어 503 으로 내려보냅니다.
    try:
        result = await run_rag(
            collection_name=qdrant_collection,
            query=req.query,
            filters=filters_dict,
            viewer_acl=current_user.acl,
            collection_names=qdrant_names,
        )
    except RuntimeError as e:
        available = llm_router.status().get("available") or []
        detail = (
            "LLM 프로바이더를 사용할 수 없습니다. "
            + ("모든 프로바이더의 Circuit Breaker 가 열려 있습니다. "
               if available else "설정된 API 키가 없습니다 (XAI_API_KEY 등을 확인하세요). ")
            + f"원인: {e}"
        )
        raise HTTPException(status_code=503, detail=detail)

    # 사용량 기록
    usage = result.get("usage", {})
    prompt_tokens     = usage.get("prompt_tokens")
    completion_tokens = usage.get("completion_tokens")
    total_tokens      = usage.get("total_tokens")
    cost = 0.0
    if prompt_tokens and completion_tokens:
        cost = (
            (prompt_tokens / 1_000_000) * COST_PER_1M_INPUT
            + (completion_tokens / 1_000_000) * COST_PER_1M_OUTPUT
        )

    try:
        session.add(UsageEvent(
            endpoint="/chat",
            model=LLM_MODEL,
            collection_id=db_collection.id if db_collection else None,
            prompt_tokens=prompt_tokens,
            completion_tokens=completion_tokens,
            total_tokens=total_tokens,
            cost_usd=cost,
            latency_ms=result.get("latency_ms"),
            cached=False,
        ))
        await session.commit()
    except Exception as e:
        print(f"Warning: 사용량 기록 실패: {e}")

    cache_set(qdrant_collection, LLM_MODEL, req.query, filters_dict, result, viewer_acl)

    return ChatResponse(
        request_id=request_id,
        answer=result["answer"],
        citations=result.get("citations", []),
        cached=False,
        latency_ms=int((time.time() - t0) * 1000),
        retrieval=result.get("retrieval", ""),
    )


# ──────────────────────────────────────────────
# Agent Endpoint (CrewAI) — NEW
# ──────────────────────────────────────────────
@app.post("/agent", response_model=AgentResponse)
async def agent_chat(
    req: AgentRequest,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """
    CrewAI 멀티 에이전트 RAG 엔드포인트.

    mode:
    - "single"        : 단일 에이전트 (빠름)
    - "rag_crew"      : 리서처 + 분석가 2-에이전트 (기본값, 균형)
    - "research_crew" : 리서처 + 팩트체커 + 작성자 3-에이전트 (심층 분석)
    """
    db_collection = await session.get(Collection, req.collection_id)
    if not db_collection:
        raise HTTPException(status_code=404, detail="지정한 컬렉션을 찾을 수 없습니다.")

    qdrant_collection = db_collection.xai_id
    request_id = str(uuid.uuid4())
    t0 = time.time()

    try:
        from agents.crew import run_single_agent, run_rag_crew, run_research_crew

        # crew.kickoff() 는 동기 호출이라 실행 중인 이벤트 루프 안에서 직접 부르면
        # CrewAI 가 거부합니다(1.x). 워커 스레드로 넘겨 실행하면 그 스레드에는
        # 이벤트 루프가 없어 정상 동작하고, 동시에 API 이벤트 루프도 막히지 않습니다.
        if req.mode == "single":
            result = await asyncio.to_thread(run_single_agent, req.query, qdrant_collection)
        elif req.mode == "research_crew":
            result = await asyncio.to_thread(run_research_crew, req.query, qdrant_collection)
        else:  # rag_crew (default)
            result = await asyncio.to_thread(run_rag_crew, req.query, qdrant_collection)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"에이전트 실행 실패: {str(e)}")

    latency_ms = int((time.time() - t0) * 1000)
    return AgentResponse(
        request_id=request_id,
        answer=result.get("answer", ""),
        mode=result.get("mode", req.mode),
        latency_ms=latency_ms,
    )
